from docx import Document
import os
import tempfile
import json

from services.addons.direct import apply_direct_redaction  # Uses apply_direct_redaction
from services.img import process_image  # Uses process_image

def process_docx_file(docx_file, redact_ocr, redact_meta, redact_face, redact_license_plate, redact_signature, redact_nsfw, is_document):
    """
    Process a .docx file by redacting text and encrypting images. Also logs detection info in a JSON file.
    """
    doc = Document(docx_file)
    upload_folder = 'static/uploads'
    detection_log = []

    # 1. Fixing TEXT Redaction
    for para in doc.paragraphs:
        original_text = para.text.strip()
        if original_text:
            redacted_text = apply_direct_redaction(original_text)  # No sensitivity
            para.clear()
            para.add_run(redacted_text)

    # 2. Handling IMAGE Redaction
    image_replacements = {}

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_filename = os.path.join(upload_folder, os.path.basename(rel.target_ref))

            # Save the original image temporarily
            with open(image_filename, "wb") as f:
                f.write(image_data)

            # Process the image using img.py
            processed_image, region_info = process_image(
                image_filename, redact_ocr, redact_meta, redact_face, 
                redact_license_plate, redact_signature, redact_nsfw, 
                is_document
            )

            # Store the processed image for later replacement
            with open(processed_image, "rb") as img_file:
                image_replacements[rel.target_ref] = img_file.read()

            # Add detection info to log
            for region in region_info:
                detection_log.append({
                    "image": rel.target_ref,
                    "region": region["region"],
                    "category": region["category"]
                })

    # Replace original images with processed images
    for rel in doc.part.rels.values():
        if rel.target_ref in image_replacements:
            rel.target_part._blob = image_replacements[rel.target_ref]

    # 3. Save the Redacted Document
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)

    # 4. Write detection log JSON in static/uploads
    json_filename = os.path.splitext(os.path.basename(docx_file))[0] + '.json'
    json_path = os.path.join('static', 'uploads', json_filename)
    with open(json_path, 'w', encoding='utf-8') as jf:
        json.dump(detection_log, jf, indent=2)

    return temp_file.name

def decrypt_docx_file(docx_path, json_path, secret_key):
    """
    Decrypts a redacted DOCX file using the provided JSON log and secret key.
    Returns the path to the decrypted DOCX file.
    """
    import json
    from services.img import decrypt_region
    from docx import Document
    import tempfile
    import os

    doc = Document(docx_path)
    upload_folder = os.path.dirname(docx_path)

    # Load JSON log
    with open(json_path, 'r', encoding='utf-8') as jf:
        detection_log = json.load(jf)

    # Group regions by image
    from collections import defaultdict
    image_regions = defaultdict(list)
    for entry in detection_log:
        image_regions[entry['image']].append(entry['region'])

    # Decrypt images
    image_replacements = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_filename = os.path.join(upload_folder, os.path.basename(rel.target_ref))
            with open(image_filename, "wb") as f:
                f.write(image_data)
            # If this image has regions to decrypt
            if rel.target_ref in image_regions:
                temp_path = image_filename
                regions = image_regions[rel.target_ref]
                for idx, region in enumerate(regions):
                    x = int(region['x'])
                    y = int(region['y'])
                    w = int(region['width'])
                    h = int(region['height'])
                    out_path = image_filename if idx == len(regions) - 1 else temp_path + f'_dec_{idx}.png'
                    decrypt_region(temp_path, x, y, w, h, secret_key, out_path)
                    temp_path = out_path
                # Read the decrypted image
                with open(temp_path, "rb") as img_file:
                    image_replacements[rel.target_ref] = img_file.read()
            else:
                # No regions to decrypt, use original
                with open(image_filename, "rb") as img_file:
                    image_replacements[rel.target_ref] = img_file.read()

    # Replace images in DOCX
    for rel in doc.part.rels.values():
        if rel.target_ref in image_replacements:
            rel.target_part._blob = image_replacements[rel.target_ref]

    # Save decrypted DOCX
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name
